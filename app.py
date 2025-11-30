import streamlit as st
import pandas as pd
import random
import os
import re
import base64
from datetime import datetime
from io import BytesIO, StringIO
from groq import Groq

# --- 0. DÉPENDANCES & SÉCURITÉ ---
try:
    from docx import Document
    from docx.shared import RGBColor
except ImportError:
    st.error("⚠️ ERREUR CRITIQUE : Le module 'python-docx' manque.")
    st.stop()

try:
    from gtts import gTTS
    HAS_AUDIO = True
except ImportError:
    HAS_AUDIO = False

# --- 1. CONFIGURATION DE LA PAGE ---
PAGE_ICON = "logo_agora.png" if os.path.exists("logo_agora.png") else "🎓"

st.set_page_config(
    page_title="Restitution PFMP - Pro'AGOrA",
    page_icon=PAGE_ICON,
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 2. STYLE CSS ---
st.markdown("""
<style>
    html, body, [class*="css"] {
        font-family: 'Segoe UI', 'Roboto', Helvetica, Arial, sans-serif;
        font-size: 16px;
    }
    .sidebar-alert {
        padding: 1rem;
        background-color: #ffebee;
        border: 1px solid #ffcdd2;
        color: #c62828;
        border-radius: 5px;
        font-weight: bold;
        font-size: 0.9rem;
        margin-bottom: 1rem;
        text-align: center;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .fixed-footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #f8f9fa;
        color: #555;
        text-align: center;
        padding: 8px 10px;
        font-size: 12px;
        border-top: 1px solid #e1e4e8;
        z-index: 99999;
    }
    [data-testid="stBottom"] {
        bottom: 50px !important;
        padding-bottom: 0px !important;
    }
    header {background-color: transparent !important;}
</style>
""", unsafe_allow_html=True)

# --- 3. GESTION ÉTAT (SESSION STATE) ---
if "messages" not in st.session_state:
    st.session_state.messages = []
if "xp" not in st.session_state:
    st.session_state.xp = 0
if "grade" not in st.session_state:
    st.session_state.grade = "👶 Stagiaire"
if "final_feedback" not in st.session_state:
    st.session_state.final_feedback = None

# SYSTÈME DE GRADES
GRADES = {
    0: "👶 Stagiaire",
    100: "👦 Assistant(e) Junior",
    300: "👨‍💼 Assistant(e) Confirmé(e)",
    600: "👩‍💻 Responsable de Pôle",
    1000: "👑 Directeur(trice)"
}

def update_xp(amount: int):
    st.session_state.xp += amount
    current_grade = "👶 Stagiaire"
    for palier, titre in GRADES.items():
        if st.session_state.xp >= palier:
            current_grade = titre
    
    if current_grade != st.session_state.grade:
        st.session_state.grade = current_grade
        st.toast(f"PROMOTION ! Tu es maintenant {current_grade} !", icon="🎉")
        st.balloons()
    else:
        st.toast(f"+{amount} XP", icon="⭐")

# --- 4. DIFFÉRENCIATION PÉDAGOGIQUE ---
PROFILES = {
    "Soutien / DYS": """
    ADAPTATION : L'élève a des difficultés ou est DYS.
    - Fais des phrases très courtes.
    - Mets les mots-clés en **gras**.
    - Utilise systématiquement des listes à puces.
    - Sois très encourageant.
    - Décompose chaque consigne en étape 1, 2, 3.
    """,
    "Standard": """
    ADAPTATION : L'élève a un niveau standard.
    - Utilise un ton professionnel bienveillant.
    - Guide l'élève s'il bloque.
    - Utilise le vocabulaire professionnel.
    """,
    "Expert / Autonomie": """
    ADAPTATION : L'élève est performant.
    - Sois exigeant.
    - Ne donne pas la réponse, pose des questions pointues ("Quel outil as-tu utilisé pour... ?").
    - Exige un vocabulaire technique précis.
    - Pousse-le à critiquer sa propre pratique.
    """
}

# --- 5. OUTILS IA (GROQ) ---
def get_api_keys_list():
    if "groq_keys" in st.secrets:
        return st.secrets["groq_keys"]
    elif "GROQ_API_KEY" in st.secrets:
        return [st.secrets["GROQ_API_KEY"]]
    return []

def query_groq_with_rotation(messages):
    available_keys = get_api_keys_list()
    if not available_keys:
        return None, "ERREUR CONFIG : Aucune clé API trouvée."
    
    keys = list(available_keys)
    random.shuffle(keys)
    models = ["llama-3.3-70b-versatile", "mixtral-8x7b-32768", "llama-3.1-8b-instant"]

    for key in keys:
        try:
            client = Groq(api_key=key)
            for model in models:
                try:
                    chat = client.chat.completions.create(
                        messages=messages,
                        model=model,
                        temperature=0.5,
                        max_tokens=1024,
                    )
                    return chat.choices[0].message.content, model
                except: continue
        except: continue
    return None, "SATURATION SERVICE."

# --- 6. OUTILS FICHIERS ---

def extract_text_from_file(uploaded_file) -> str:
    try:
        filename = uploaded_file.name.lower()
        if filename.endswith(".docx"):
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])[:15000]
        elif filename.endswith(".xlsx") or filename.endswith(".xls"):
            df = pd.read_excel(uploaded_file)
            return df.to_string()[:15000]
        elif filename.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
            return df.to_string()[:15000]
        else:
            return "Format non supporté."
    except Exception as e:
        return f"Erreur de lecture : {str(e)}"

def create_docx_history(messages, student_name, final_feedback=None):
    doc = Document()
    doc.add_heading(f"Restitution PFMP - {student_name}", 0)
    doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    if final_feedback:
        doc.add_heading("BILAN PÉDAGOGIQUE (IA)", level=1)
        doc.add_paragraph(final_feedback)
        doc.add_page_break()

    doc.add_heading("Historique de la conversation", level=1)
    for msg in messages:
        if msg["role"] == "system": continue
        role_name = "SUPERVISEUR (IA)" if msg["role"] == "assistant" else student_name.upper()
        p = doc.add_paragraph()
        runner = p.add_run(f"{role_name} :")
        runner.bold = True
        if msg["role"] == "assistant":
            runner.font.color.rgb = RGBColor(0, 0, 255)
        else:
            runner.font.color.rgb = RGBColor(0, 100, 0)
        doc.add_paragraph(msg["content"])
        doc.add_paragraph("") 
        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def clean_text_for_audio(text: str) -> str:
    text = re.sub(r"[\*_]{1,3}", "", text)
    text = re.sub(r"\[.*?\]", "", text)
    return text[:500]

# --- 7. PROMPT SYSTÈME DYNAMIQUE ---
def get_system_prompt(profile_key):
    differentiation_instruction = PROFILES.get(profile_key, PROFILES["Standard"])
    
    base_prompt = f"""
RÔLE : Tu es le Superviseur Professionnel de l'élève (Bac Pro AGOrA).
OBJECTIF : Aider l'élève à analyser son vécu (PFMP) selon la méthode de l'EXPLICITATION (Vermersch).
Tu ne cherches pas le "bon résultat", mais à comprendre le CHEMINEMENT de l'élève.

{differentiation_instruction}

CONSIGNES PÉDAGOGIQUES (EXPLICITATION) :
1. Questionne sur le "COMMENT" : "Par quoi as-tu commencé ?", "Qu'as-tu fait juste après ?", "Comment savais-tu que... ?".
2. Si l'élève est vague ("J'ai fait un courrier"), demande des détails procéduraux ("Quel logiciel ?", "A partir de quel modèle ?", "Qui t'a donné les infos ?").
3. SÉCURITÉ : Stoppe immédiatement toute donnée personnelle réelle.
"""
    return base_prompt

INITIAL_MESSAGE = """
👋 **Bonjour.**

Je suis ton Superviseur Virtuel OECAM. Nous allons faire le point sur ton stage.

Raconte-moi une activité importante que tu as réalisée. Essaie d'être précis sur **comment** tu as fait.
"""

if not st.session_state.messages:
    st.session_state.messages.append({"role": "assistant", "content": INITIAL_MESSAGE})

# --- 8. INTERFACE GRAPHIQUE ---

st.title("🎓 Restitution PFMP & Analyse de Pratique")

with st.sidebar:
    if os.path.exists(PAGE_ICON):
        st.image(PAGE_ICON, width=80)
    
    st.markdown(f"### 🏆 {st.session_state.grade}")
    st.progress(min(st.session_state.xp / 1000, 1.0))
    student_name = st.text_input("Ton Prénom :", placeholder="Ex: Thomas")
    
    # --- SÉLECTEUR DE DIFFÉRENCIATION ---
    st.markdown("### ⚙️ Niveau d'aide")
    selected_profile = st.selectbox(
        "Choisis ton profil :",
        ["Soutien / DYS", "Standard", "Expert / Autonomie"],
        index=1
    )

    st.divider()

    st.markdown("""
    <div class="sidebar-alert">
    🚫 <b>ANONYMISE TOUT !</b><br>
    Pas de vrais noms, pas de vrais téléphones.
    </div>
    """, unsafe_allow_html=True)

    st.divider()

    st.subheader("📂 Analyser un document")
    uploaded_file = st.file_uploader("Rapport/Brouillon", type=['docx', 'xlsx', 'xls', 'csv'])
    
    if uploaded_file and student_name:
        if st.button("🚀 Envoyer à l'analyse"):
            with st.spinner("Lecture..."):
                text = extract_text_from_file(uploaded_file)
                st.session_state.messages.append({"role": "user", "content": f"Voici mon document ({uploaded_file.name}) :\n\n{text}"})
                update_xp(50)
                st.rerun()

    st.divider()

    # --- BOUTON DE FIN (FEEDBACK) ---
    st.subheader("🏁 Fin de séance")
    if st.button("Générer le Bilan Pédagogique", type="primary"):
        if len(st.session_state.messages) < 4:
            st.warning("Discute un peu plus avant de générer le bilan.")
        else:
            with st.spinner("Rédaction du bilan d'explicitation..."):
                # Prompt spécifique pour le bilan final
                history_txt = "\n".join([m['content'] for m in st.session_state.messages])
                prompt_bilan = f"""
                Agis comme un expert en pédagogie. Analyse cette conversation d'explicitation avec un élève de Bac Pro.
                
                CONVERSATION :
                {history_txt}
                
                Génère un bilan structuré adressé à l'élève (utilise le "TU") :
                1. **Contexte identifié** : Où et quoi ?
                2. **Procédure décrite** : L'élève a-t-il su expliquer "comment" il a fait ? (Oui/Non/Partiellement).
                3. **Points forts détectés** : Vocabulaire, outils, posture pro.
                4. **Conseil d'amélioration** : Sur quoi progresser pour le CCF.
                
                Sois bienveillant et constructif.
                """
                response_bilan, _ = query_groq_with_rotation([{"role": "user", "content": prompt_bilan}])
                st.session_state.final_feedback = response_bilan
                st.rerun()

    # --- EXPORT ---
    if st.session_state.final_feedback:
        st.success("Bilan généré ! (Voir en bas de page)")
        docx_file = create_docx_history(st.session_state.messages, student_name, st.session_state.final_feedback)
        st.download_button(
            "📄 Télécharger Bilan + Chat (.docx)",
            data=docx_file,
            file_name=f"Bilan_PFMP_{student_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    if st.button("🗑️ Reset"):
        st.session_state.messages = [{"role": "assistant", "content": INITIAL_MESSAGE}]
        st.session_state.xp = 0
        st.session_state.final_feedback = None
        st.rerun()

# CHAT
chat_container = st.container()
with chat_container:
    for msg in st.session_state.messages:
        role_avatar = "🤖" if msg["role"] == "assistant" else "🧑‍🎓"
        with st.chat_message(msg["role"], avatar=role_avatar):
            if "Voici mon document" in msg["content"]:
                st.info("📄 *Document transmis.*")
            else:
                st.markdown(msg["content"])
                if msg["role"] == "assistant" and HAS_AUDIO:
                    if st.button("🔊", key=str(random.randint(0,100000))): # Clé unique simple
                        try:
                            tts = gTTS(clean_text_for_audio(msg["content"]), lang="fr")
                            buf = BytesIO()
                            tts.write_to_fp(buf)
                            st.audio(buf, format="audio/mp3", start_time=0)
                        except: pass

# AFFICHAGE DU FEEDBACK FINAL
if st.session_state.final_feedback:
    st.markdown("---")
    st.markdown("## 📝 Bilan de l'entretien d'explicitation")
    st.info(st.session_state.final_feedback)

st.write("<br><br>", unsafe_allow_html=True)

# FOOTER
st.markdown("""
<div class="fixed-footer">
    ℹ️ <b>IA Pédagogique</b> - Différenciation activée. Ne pas saisir de données réelles.
</div>
""", unsafe_allow_html=True)

# INPUT
if user_input := st.chat_input("Réponds au superviseur..."):
    if not student_name:
        st.toast("⚠️ Indique ton prénom !", icon="👉")
    else:
        st.session_state.messages.append({"role": "user", "content": user_input})
        update_xp(10)
        
        with st.chat_message("assistant", avatar="🤖"):
            with st.spinner("Analyse en cours..."):
                # On utilise le prompt dynamique selon le profil choisi
                current_system_prompt = get_system_prompt(selected_profile)
                
                messages_payload = [{"role": "system", "content": current_system_prompt}]
                messages_payload.extend(st.session_state.messages[-8:])
                
                response_content, _ = query_groq_with_rotation(messages_payload)
                if not response_content: response_content = "⚠️ Erreur IA."
                st.markdown(response_content)
        
        st.session_state.messages.append({"role": "assistant", "content": response_content})
        st.rerun()
